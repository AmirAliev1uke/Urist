"""Универсальный парсер документов.

Поддерживаемые форматы (для базы знаний и для анализа):
  - PDF  (через PyMuPDF)
  - DOCX (через python-docx)
  - DOC  (конвертация в DOCX через LibreOffice headless)

DOC — старый бинарный формат Word 97-2003. python-docx его не читает, поэтому
конвертируем в DOCX через LibreOffice. Это надёжнее, чем извлекать «грязный»
текст напрямую, и сохраняет таблицы.
"""
import io
import os
import shutil
import subprocess
import tempfile
from dataclasses import dataclass

from app.core.pdf_parser import parse_pdf


@dataclass
class ParsedDocument:
    """Результат извлечения текста из документа любого формата."""

    file_name: str
    file_type: str  # 'pdf' | 'docx' | 'doc'
    pages: list[str]  # текст постранично/поабзацно
    total_chars: int

    @property
    def full_text(self) -> str:
        return "\n\n".join(p for p in self.pages if p.strip())


def parse_document(file_bytes: bytes, file_name: str = "document") -> ParsedDocument:
    """Определить формат по имени файла и распарсить соответствующим парсером.

    Бросает ValueError для неподдерживаемых форматов или пустых документов.
    """
    name = file_name.lower()

    if name.endswith(".pdf"):
        parsed = parse_pdf(file_bytes, file_name=file_name)
        return ParsedDocument(
            file_name=parsed.file_name,
            file_type="pdf",
            pages=[p.text for p in parsed.pages],
            total_chars=parsed.total_chars,
        )

    if name.endswith(".docx"):
        return _parse_docx(file_bytes, file_name)

    if name.endswith(".doc"):
        return _parse_doc(file_bytes, file_name)

    raise ValueError(
        f"Неподдерживаемый формат файла «{file_name}». "
        "Доступные форматы: PDF, DOCX, DOC."
    )


def _parse_docx(file_bytes: bytes, file_name: str) -> ParsedDocument:
    """Извлечь текст из DOCX через python-docx."""
    try:
        from docx import Document  # noqa: WPS433 — ленивый импорт
    except ImportError as exc:
        raise RuntimeError("Библиотека python-docx не установлена.") from exc

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"Не удалось прочитать DOCX: {exc}") from exc

    # Текст из абзацев + из таблиц (договоры часто табличные)
    paragraphs: list[str] = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    if not paragraphs:
        raise ValueError(
            "DOCX не содержит извлекаемого текста. "
            "Возможно, документ состоит только из изображений."
        )

    return ParsedDocument(
        file_name=file_name,
        file_type="docx",
        pages=paragraphs,
        total_chars=sum(len(p) for p in paragraphs),
    )


def _parse_doc(file_bytes: bytes, file_name: str) -> ParsedDocument:
    """Извлечь текст из DOC через конвертацию в DOCX (LibreOffice headless).

    DOC (Word 97-2003) — бинарный формат. python-docx его не поддерживает.
    Конвертируем в DOCX, затем парсим как DOCX. Это сохраняет таблицы и кириллицу.
    Требует установленного libreoffice в образе (см. Dockerfile).
    """
    if not shutil.which("soffice") and not shutil.which("libreoffice"):
        raise RuntimeError(
            "Для обработки .doc файлов требуется LibreOffice (soffice). "
            "Он не найден в системе. Обратитесь к администратору."
        )

    # Пишем .doc во временную директорию
    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = os.path.join(tmpdir, "input.doc")
        with open(doc_path, "wb") as f:
            f.write(file_bytes)

        # Конвертируем в DOCX: soffice --headless --convert-to docx <file>
        cmd = [
            "soffice",
            "--headless",
            "--norestore",
            "--convert-to", "docx",
            "--outdir", tmpdir,
            doc_path,
        ]
        try:
            result = subprocess.run(
                cmd, capture_output=True, timeout=120, check=False
            )
        except subprocess.TimeoutExpired as exc:
            raise ValueError(
                "Превышено время конвертации .doc файла (120с). "
                "Возможно, файл повреждён."
            ) from exc

        converted_path = os.path.join(tmpdir, "input.docx")
        if not os.path.exists(converted_path):
            stderr = result.stderr.decode("utf-8", errors="replace") if result.stderr else ""
            raise ValueError(
                f"Не удалось конвертировать .doc в .docx. "
                f"Возможно, файл повреждён или это не DOC. {stderr}"
            )

        with open(converted_path, "rb") as f:
            docx_bytes = f.read()

    parsed = _parse_docx(docx_bytes, file_name)
    return ParsedDocument(
        file_name=parsed.file_name,
        file_type="doc",  # сохраняем исходный тип для логов
        pages=parsed.pages,
        total_chars=parsed.total_chars,
    )
